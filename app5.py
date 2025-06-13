import io
import sqlite3
from datetime import datetime
from pprint import pprint
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import qrcode
import streamlit as st
from PIL import Image
from scipy.integrate import quad
from scipy.optimize import differential_evolution
from scipy.stats import poisson
from docx import Document
from docx.shared import Inches

# Initialize session state to store selected templates
if "selected_template" not in st.session_state:
    st.session_state.selected_template = "Preventive Maintenance"
if "selected_inventory_part" not in st.session_state:
    st.session_state.selected_inventory_part = None
if "selected_technician" not in st.session_state:
    st.session_state.selected_technician = None
if "form_active" not in st.session_state:
    st.session_state.form_active = False

# Database setup
def init_db():
    print("connecting to database ...")
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    
    # Create work_orders table
    c.execute('''CREATE TABLE IF NOT EXISTS work_orders (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 template_type TEXT,
                 asset_id TEXT,
                 description TEXT,
                 priority TEXT,
                 requested_date TEXT
                 )''')
    
    # Create inventory table
    c.execute('''CREATE TABLE IF NOT EXISTS inventory (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 part_id TEXT,
                 name_description TEXT,
                 location TEXT,
                 quantity_on_hand INTEGER,
                 min_level INTEGER,
                 max_level INTEGER,
                 last_restock_date TEXT,
                 supplier_info TEXT
                 )''')
    
    # Create technicians table
    c.execute('''CREATE TABLE IF NOT EXISTS technicians (
                 id INTEGER PRIMARY KEY AUTOINCREMENT,
                 name TEXT,
                 technician_id TEXT,
                 contact_details TEXT,
                 certifications TEXT,
                 skill_sets TEXT,
                 experience_level TEXT,
                 work_location TEXT,
                 shift_schedule TEXT
                 )''')
    
    conn.commit()
    conn.close()
    print("database initialized\n")

# Initialize database
init_db()

# Load work orders from database (for QR code and Word document)
def load_work_orders():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM work_orders")
    work_orders = []
    for row in c.fetchall():
        work_orders.append({
            "id": row[0],
            "Template Type": row[1],
            "Asset ID": row[2],
            "Description": row[3],
            "Priority": row[4],
            "Requested Date": row[5]
        })
    conn.close()
    return work_orders

# Load inventory from database (for QR code and Word document)
def load_inventory():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM inventory")
    inventory = []
    for row in c.fetchall():
        inventory.append({
            "id": row[0],
            "Part ID/SKU": row[1],
            "Name and Description": row[2],
            "Location": row[3],
            "Quantity on Hand": row[4],
            "Min Level": row[5],
            "Max Level": row[6],
            "Last Restock Date": row[7],
            "Supplier Info": row[8]
        })
    conn.close()
    return inventory

# Load technicians from database (for QR code and Word document)
def load_technicians():
    conn = sqlite3.connect("reliability_data.db")
    c = conn.cursor()
    c.execute("SELECT * FROM technicians")
    technicians = []
    for row in c.fetchall():
        technicians.append({
            "id": row[0],
            "Name": row[1],
            "Technician ID": row[2],
            "Contact Details": row[3],
            "Certifications": row[4],
            "Skill Sets": row[5],
            "Experience Level": row[6],
            "Work Location": row[7],
            "Shift Schedule": row[8]
        })
    conn.close()
    return technicians

# Function to generate Word document
def generate_word_document():
    work_orders = load_work_orders()
    inventory = load_inventory()
    technicians = load_technicians()
    
    doc = Document()
    doc.add_heading('MISSION RELIABILITY EVALUATOR - Saved Data', 0)
    
    # Work Orders Section
    doc.add_heading('Work Orders', level=1)
    if work_orders:
        for idx, order in enumerate(work_orders):
            doc.add_heading(f'Work Order {idx + 1} - {order["Template Type"]}', level=2)
            doc.add_paragraph(f"Template Type: {order['Template Type']}")
            doc.add_paragraph(f"Asset ID: {order['Asset ID']}")
            doc.add_paragraph(f"Description: {order['Description']}")
            doc.add_paragraph(f"Priority: {order['Priority']}")
            doc.add_paragraph(f"Requested Date: {order['Requested Date']}")
            doc.add_paragraph()  # Add spacing
    else:
        doc.add_paragraph("No work orders saved.")
    
    # Inventory Section
    doc.add_heading('Inventory', level=1)
    if inventory:
        for idx, item in enumerate(inventory):
            doc.add_heading(f'Part {idx + 1} - {item["Part ID/SKU"]}', level=2)
            doc.add_paragraph(f"Part ID/SKU: {item['Part ID/SKU']}")
            doc.add_paragraph(f"Name and Description: {item['Name and Description']}")
            doc.add_paragraph(f"Location: {item['Location']}")
            doc.add_paragraph(f"Quantity on Hand: {item['Quantity on Hand']}")
            doc.add_paragraph(f"Min Level: {item['Min Level']}")
            doc.add_paragraph(f"Max Level: {item['Max Level']}")
            doc.add_paragraph(f"Last Restock Date: {item['Last Restock Date']}")
            doc.add_paragraph(f"Supplier Info: {item['Supplier Info']}")
            if item['Quantity on Hand'] < item['Min Level']:
                doc.add_paragraph("Warning: Quantity on Hand is below the Minimum Level!", style='Intense Quote')
            doc.add_paragraph()
    else:
        doc.add_paragraph("No inventory items saved.")
    
    # Technicians Section
    doc.add_heading('Technician Profiles', level=1)
    if technicians:
        for idx, tech in enumerate(technicians):
            doc.add_heading(f'Technician {idx + 1} - {tech["Name"]}', level=2)
            doc.add_paragraph(f"Name: {tech['Name']}")
            doc.add_paragraph(f"Technician ID: {tech['Technician ID']}")
            doc.add_paragraph(f"Contact Details: {tech['Contact Details']}")
            doc.add_paragraph(f"Certifications & Licenses: {tech['Certifications']}")
            doc.add_paragraph(f"Skill Sets: {tech['Skill Sets']}")
            doc.add_paragraph(f"Experience Level: {tech['Experience Level']}")
            doc.add_paragraph(f"Work Location/Zone: {tech['Work Location']}")
            doc.add_paragraph(f"Shift Schedule and Availability: {tech['Shift Schedule']}")
            doc.add_paragraph()
    else:
        doc.add_paragraph("No technician profiles saved.")
    
    # Save document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App Layout
st.title("MISSION RELIABILITY EVALUATOR")

# Subscription Details Section
st.header("Subscription Details")
st.write("Streamline managing work orders, PMs, & assets. Reduce downtime & costs.")
st.write("Only $35/month per user")

# List of features
features = {
    "Work Order Management": "Create, track, and manage work orders efficiently to ensure timely completion of tasks. Includes editable templates for: Preventive maintenance, Reactive maintenance, Inspections. Fields typically include: Asset/equipment ID, Description of issue/task, Priority level, Requested date/time.",
    "Asset & Equipment Management": "Monitor and manage all assets and equipment, including their maintenance history and performance.",
    "Preventive Maintenance": "Schedule and automate preventive maintenance tasks to reduce unexpected breakdowns.",
    "Vendor Management": "Manage vendor information, contracts, and communications to streamline external support.",
    "Maintenance Requests": "Allow staff to submit maintenance requests easily and track their status in real-time.",
    "Parts & Inventory Management": "Track inventory levels, manage spare parts, and reorder supplies to avoid delays.",
    "Resource & Labor Management": "Assign resources and labor to tasks, ensuring optimal workforce utilization.",
    "Reports & KPIs": "Generate detailed reports and key performance indicators to monitor system performance.",
    "Web & Mobile App": "Access the system via web or mobile app for on-the-go management and updates.",
    "QR Code Scanning": "Use QR codes to quickly access asset details, work orders, or maintenance history.",
    "Lifetime Customer Support": "Get unlimited support from our team to ensure smooth operation of your system.",
    "Downtime Tracking": "Monitor and analyze equipment downtime to identify patterns and reduce future occurrences.",
    "Cost Analysis": "Track maintenance costs and analyze expenses to optimize budget allocation.",
    "Custom Notifications": "Set up custom alerts for upcoming maintenance, overdue tasks, or low inventory levels.",
    "Predictive Maintenance": "Leverage AI to predict equipment failures before they occur, minimizing downtime.",
    "AI-Driven Insights": "Gain actionable insights from data analysis to optimize maintenance strategies.",
    "Automated Scheduling": "Automatically schedule maintenance tasks based on equipment usage and priority.",
    "Technician Profiles & Skills Tracking": "Track detailed info about each technician, including skills, certifications, and availability."
}

# Work Order Management Section
if st.button("Work Order Management") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) work order button clicked")
    st.write("**Work Order Management**: Create, track, and manage work orders efficiently to ensure timely completion of tasks. Includes editable templates for: Preventive maintenance, Reactive maintenance, Inspections. Fields typically include: Asset/equipment ID, Description of issue/task, Priority level, Requested date/time.")
    st.subheader("Create Work Order")

    # Form for creating work orders
    with st.form(key="work_order_form"):
        print("Creating work order form...")
        selected_template = st.selectbox("Select Template Type", ["Preventive Maintenance", "Reactive Maintenance", "Inspections"], index=["Preventive Maintenance", "Reactive Maintenance", "Inspections"].index(st.session_state.selected_template), key="template_type")
        asset_id = st.text_input("Asset/Equipment ID", value=st.session_state.get("asset_id", "EQ001"), key="asset_id_input")
        description = st.text_area("Description of Issue/Task", value=st.session_state.get("description", "Routine check-up"), key="description_input")
        priority = st.selectbox("Priority Level", ["Low", "Medium", "High"], index=["Low", "Medium", "High"].index(st.session_state.get("priority", "Medium")), key="priority_input")
        requested_date = st.text_input("Requested Date/Time (YYYY-MM-DD HH:MM)", value=st.session_state.get("requested_date", "2025-06-04 09:00"), key="requested_date_input")
        submit_button = st.form_submit_button(label="Save Work Order")

        if submit_button:
            print("(BTN) The submit button was clicked")
            # Validate date
            try:
                datetime.strptime(requested_date, "%Y-%m-%d %H:%M")
            except ValueError:
                st.error("Please enter the Requested Date/Time in YYYY-MM-DD HH:MM format.")
                st.stop()

            # Create work order
            work_order = {
                "Template Type": selected_template,
                "Asset ID": asset_id,
                "Description": description,
                "Priority": priority,
                "Requested Date": requested_date
            }
            
            pprint(work_order)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            c.execute('''INSERT INTO work_orders (template_type, asset_id, description, priority, requested_date) VALUES (?, ?, ?, ?, ?)''',
                      (selected_template, asset_id, description, priority, requested_date))
            conn.commit()
            conn.close()
            st.session_state.selected_template = selected_template
            st.session_state["asset_id"] = asset_id
            st.session_state["description"] = description
            st.session_state["priority"] = priority
            st.session_state["requested_date"] = requested_date
            st.success("Work order saved successfully! Access saved data via QR code or Word document.")

# Parts & Inventory Management Section
if st.button("Parts & Inventory Management") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) parts & inventory management button clicked")
    st.write("**Parts & Inventory Management**: Track inventory levels, manage spare parts, and reorder supplies to avoid delays.")
    st.subheader("Add Inventory Item")

    # Form for adding inventory items
    with st.form(key="inventory_form"):
        print("Creating inventory form...")
        part_id = st.text_input("Part ID/SKU", value=st.session_state.get("part_id", "PART001"), key="part_id_input")
        name_description = st.text_area("Name and Description", value=st.session_state.get("name_description", "Oil Filter - Standard Size"), key="name_description")
        location = st.text_input("Location (warehouse, site, truck, etc.)", value=st.session_state.get("location", "Warehouse A"), key="location_input")
        quantity_on_hand = st.number_input("Quantity on Hand", min_value=0, value=st.session_state.get("quantity_on_hand", 10), key="quantity_on_hand_input")
        min_level = st.number_input("Minimum Level", min_value=0, value=st.session_state.get("min_level", 5), key="min_level")
        max_level = st.number_input("Maximum Level", min_value=0, value=st.session_state.get("max_level", 50), key="maximum_level")
        last_restock_date = st.text_input("Last Restock Date (YYYY-MM-DD)", value=st.session_state.get("last_restock_date", "2025-06-05"), key="restock_date")
        supplier_info = st.text_area("Supplier/Vendor Information", value=st.session_state.get("supplier_info", "Supplier: ABC Corp\nContact: 555-1234"), key="supplier_type")
        submit_button = st.form_submit_button("label="Submit")

        if submit_button:
            print("(button) The submit button was clicked")
            # Validate inputs
            try:
                datetime.strptime(last_restock_date, "%Y-%m-%d")
            except ValueError:
                st.error("Please enter the Last Restock Date in YYYY-MM-DD format.")
                st.stop()

            # Create inventory item
            inventory_item = {
                "Part ID/SKU": part_id,
                "Name and Description": name_description,
                "Location": location,
                "Quantity on Hand": quantity_on_hand,
                "Min Level": min_level,
                "Max Level": max_level,
                "Last Restock Date": last_restock_date,
                "Supplier Info": supplier_info
            }
            pprint(inventory_item)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            c.execute('''INSERT INTO inventory (part_id, name_description, location, quantity_on_hand, min_level, max_level, last_restock_date, supplier_info) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                      (part_id, name_description, location, quantity_on_hand, min_level, max_level, last_restock_date, supplier_info))
            conn.commit()
            conn.close()
            st.session_state["part_id"] = part_id
            st.session_state["name_description"] = name_description
            st.session_state["location"] = location
            st.session_state["quantity_on_hand"] = quantity_on_hand
            st.session_state["min_level"] = min_level
            st.session_state["max_level"] = max_level
            st.session_state["last_restock_date"] = last_restock_date
            st.session_state["supplier_info"] = supplier_info
            st.success("Inventory item saved successfully! Access saved data via QR code or Word document.")

# Technician Profiles & Skills Tracking Section
if st.button("Technician Profiles & Skills Tracking") or st.session_state.get("form_active", False):
    st.session_state.form_active = True
    print("\n(BTN) technician profiles & skills tracking button clicked")
    st.write("**Technician Profiles & Skills Tracking**: Track detailed info about each technician, including skills, certifications, and availability.")
    st.subheader("Add Technician Profile")

    # Form for adding technicians
    with st.form(key="technician_form"):
        print("Creating technician form...")
        name = st.text_input("Name", value=st.session_state.get("technician_name", "John Doe"), key="technician_name_input")
        technician_id = st.text_input("Technician ID", value=st.session_state.get("technician_id", "TECH001"), key="technician_id_input")
        contact_details = st.text_area("Contact Details", value=st.session_state.get("contact_details", "Phone: 555-1234\nEmail: john.doe@example.com"), key="contact_details_input")
        certifications = st.text_area("Certifications & Licenses", value=st.session_state.get("certifications", "Certified Welder, Electrical Safety License"), key="certifications_input")
        skill_sets = st.text_area("Skill Sets (e.g., welding, electrical, corrosion assessment)", value=st.session_state.get("skill_sets", "Welding, Electrical, Corrosion Assessment"), key="skill_sets_input")
        experience_level = st.selectbox("Experience Level", ["Entry Level", "Intermediate", "Senior"], index=["Entry Level", "Intermediate", "Senior"].index(st.session_state.get("experience_level", "Intermediate")), key="experience_level_input")
        work_location = st.text_input("Work Location/Zone", value=st.session_state.get("work_location", "Zone A"), key="work_location_input")
        shift_schedule = st.text_area("Shift Schedule and Availability", value=st.session_state.get("shift_schedule", "Mon-Fri, 8 AM - 4 PM\nAvailable for overtime"), key="shift_schedule_input")
        submit_button = st.form_submit_button(label="Save Technician Profile")

        if submit_button:
            print("(BTN) The submit button was clicked")
            # Create technician profile
            technician = {
                "Name": name,
                "Technician ID": technician_id,
                "Contact Details": contact_details,
                "Certifications": certifications,
                "Skill Sets": skill_sets,
                "Experience Level": experience_level,
                "Work Location": work_location,
                "Shift Schedule": shift_schedule
            }
            pprint(technician)
            conn = sqlite3.connect("reliability_data.db")
            c = conn.cursor()
            c.execute('''INSERT INTO technicians (name, technician_id, contact_details, certifications, skill_sets, experience_level, work_location, shift_schedule) VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                      (name, technician_id, contact_details, certifications, skill_sets, experience_level, work_location, shift_schedule))
            conn.commit()
            conn.close()
            st.session_state["technician_name"] = name
            st.session_state["technician_id"] = technician_id
            st.session_state["contact_details"] = contact_details
            st.session_state["certifications"] = certifications
            st.session_state["skill_sets"] = skill_sets
            st.session_state["experience_level"] = experience_level
            st.session_state["work_location"] = work_location
            st.session_state["shift_schedule"] = shift_schedule
            st.success("Technician profile saved successfully! Access saved data via QR code or Word document.")

# Download Saved Data as Word Document
if st.button("Download Saved Data as Word Document"):
    st.header("Download All Saved Data")
    st.write("Click below to download all saved Work Orders, Inventory Items, and Technician Profiles as a Word document.")
    
    # Generate the Word document
    doc_buffer = generate_word_document()
    
    # Provide download button
    st.download_button(
        label="Download Word Document",
        data=doc_buffer,
        file_name="Reliability_Data.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# QR Code Generation Section for All Saved Data
if st.button("Generate QR Code for All Saved Data"):
    st.header("QR Code for All Saved Data")
    st.write("Scan this QR code to view all saved Work Orders, Inventory Items, and Technician Profiles. Note: Data is condensed; scan with a QR reader to see details.")

    # Load all data
    work_orders = load_work_orders()
    inventory = load_inventory()
    technicians = load_technicians()

    # Format data into a concise string
    data_lines = []
    for order in work_orders:
        data_lines.append(f"WO:{order['Template Type']},{order['Asset ID']},{order['Priority']},{order['Requested Date']}")
    for item in inventory:
        data_lines.append(f"INV:{item['Part ID/SKU']},{item['Name and Description']},{item['Quantity on Hand']},{item['Location']}")
    for tech in technicians:
        data_lines.append(f"TECH:{tech['Name']},{tech['Technician ID']},{tech['Skill Sets']},{tech['Work Location']}")
    qr_data = "\n".join(data_lines) if data_lines else "No data saved."

    # Generate QR code
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=5,
        border=2,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)

    # Create an image from the QR Code instance
    img = qr.make_image(fill_color="black", back_color="white")

    # Save the image to a bytes buffer to display in Streamlit
    img_buffer = io.BytesIO()
    img.save(img_buffer, format="PNG")
    img_buffer.seek(0)

    # Display the QR code in Streamlit with a fixed width
    st.image(img_buffer, caption="QR Code containing all saved data", width=200)

# Parameters Input
st.header("Parameters")
col1, col2 = st.columns(2)

with col1:
    W = st.number_input("Workload (W, Gops)", value=600.0)
    g = st.number_input("Processing Speed (g, Gops/hr)", value=12.0)
    lambda_ = st.number_input("Shock Rate (lambda, /hr)", value=0.25, step=0.01)
    lambda_tilde = st.number_input("Shock Rate during Rescue (lambda_tilde, /hr)", value=0.06, step=0.01)
    alpha = st.number_input("Sharing Factor (alpha)", value=0.8, step=0.1)
    w = st.number_input("Inspection Workload (w, Gops)", value=25.0)
    epsilon = st.number_input("Detection Cutoff (epsilon)", value=0.6, step=0.1)
    p = st.number_input("False Negative Rate (p)", value=0.05, step=0.01)
    S_star_85 = st.number_input("Constraint S >= 0.85 (S*)", value=0.85, step=0.01)

with col2:
    q = st.number_input("False Positive Rate (q)", value=0.03, step=0.01)
    delta = st.number_input("Deceleration during Rescue (delta)", value=0.7, step=0.1)
    mu0 = st.number_input("Base Rescue Time (mu_0, hr)", value=12.0)
    mu1 = st.number_input("Rescue Time Growth (mu_1)", value=0.1, step=0.1)
    eta = st.number_input("Weibull Scale (eta, hr)", value=120.0)
    beta = st.number_input("Weibull Shape (beta)", value=2.0)
    max_m = st.number_input("Max Shocks (m_max)", value=50, step=1, format="%d")
    N = st.number_input("Number of Inspections (N)", value=1, step=1, format="%d")
    S_star_90 = st.number_input("Constraint S >= 0.90 (S*)", value=0.90, step=0.01)

# Functions
def theta():
    return w / (alpha * g)

def total_mission_time(N):
    return (W + N * w) / g

def phi_i(tau, i):
    return mu0 + mu1 * (tau[i-1] * g + w / alpha - i * w) / W

def z(k):
    return 1 if k == 0 else 0.97 * (0.85) ** (k - 1)

def Z(m):
    return np.prod([z(i) for i in range(m + 1)])

def P(t, m, lambda_val):
    return poisson.pmf(m, lambda_val * t)

def u(t):
    return lambda_ * sum(P(t, m-1, lambda_) * (1 - z(m)) * Z(m-1) for m in range(1, max_m))

def u_tilde(t, tau_i, theta_val):
    return lambda_tilde * sum(
        P(tau_i + theta_val, k, lambda_) * 
        sum(P(t, l, lambda_tilde) * Z(k + l) for l in range(max_m))
        for k in range(max_m)
    )

def V(t):
    return 1 - np.exp(-((t / eta) ** beta)) if t > 0 else 0

def calculate_mission_success_probability(tau, N, T, theta_val):
    tau = [0] + (list(tau) if np.isscalar(tau) or len(tau) > 0 else []) + [T]
    R = 0
    if N > 0:
        R += (1 - q) ** N * sum(P(T, m, lambda_) * Z(m) for m in range(max_m))
        for i in range(1, N + 1):
            integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 
                              tau[i-1] + epsilon * theta_val, 
                              tau[i] + epsilon * theta_val)
            R += (1 - q) ** (i-1) * p ** (N - i + 1) * integral
        integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 
                          tau[N] + epsilon * theta_val, T)
        R += (1 - q) ** N * integral
    else:
        R += sum(P(T, m, lambda_) * Z(m) for m in range(max_m))
        integral, _ = quad(lambda t: (1 - V(T - t)) * u(t), 0, T)
        R += integral
    return R

def calculate_failure_avoidance_probability(tau, N, T, theta_val):
    tau = [0] + (list(tau) if np.isscalar(tau) or len(tau) > 0 else []) + [T]
    S = 0
    for i in range(1, N + 1):
        phi = phi_i(tau, i)
        for k in range(1, i + 1):
            integral, _ = quad(
                lambda t: (1 - V(tau[i] + theta_val - t + delta * phi)) * u(t),
                tau[k-1] + epsilon * theta_val, 
                tau[k] + epsilon * theta_val)
            S += (1 - q) ** (k-1) * p ** (i-k) * (1 - p) * integral
    for i in range(1, N + 1):
        phi = phi_i(tau, i)
        term1 = sum(P(tau[i] + theta_val, k, lambda_) * 
                    sum(P(phi, l, lambda_tilde) * Z(k + l) for l in range(max_m))
                    for k in range(max_m))
        integral, _ = quad(
            lambda t: (1 - V(delta * (phi - t))) * u_tilde(t, tau[i], theta_val),
            0, phi)
        S += q * (1 - q) ** (i-1) * (term1 + integral)
    S += calculate_mission_success_probability(tau[1:-1], N, T, theta_val)
    return S

def objective_1(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 0
    theta_val = theta()
    T = total_mission_time(N)
    R = calculate_mission_success_probability([], N, T, theta_val)
    S = calculate_failure_avoidance_probability([], N, T, theta_val)
    return R, S

def objective_2(N, lambda_val):
    global lambda_
    lambda_ = lambda_val
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        return -calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    S_de = -result_de.fun
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

def objective_3(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 1
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        S = calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
        R = calculate_mission_success_probability([tau[0]], N, T, theta_val)
        penalty = 1e6 * max(0, S_star_90 - S)
        return -R + penalty
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    S_de = calculate_failure_avoidance_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

def objective_4(lambda_val):
    global lambda_
    lambda_ = lambda_val
    N = 1
    theta_val = theta()
    T = total_mission_time(N)
    def objective_de(tau):
        S = calculate_failure_avoidance_probability([tau[0]], N, T, theta_val)
        R = calculate_mission_success_probability([tau[0]], N, T, theta_val)
        penalty = 1e6 * max(0, S_star_85 - S)
        return -R + penalty
    result_de = differential_evolution(objective_de, bounds=[(0, T)], maxiter=50)
    tau_de = [result_de.x[0]]
    R_de = calculate_mission_success_probability(tau_de, N, T, theta_val)
    S_de = calculate_failure_avoidance_probability(tau_de, N, T, theta_val)
    return R_de, S_de, tau_de

# Calculate Button
if st.button("Calculate"):
    st.header("Results")
    
    # Objective 1
    st.subheader("Objective 1: R = S, No Inspections")
    R, S = objective_1(lambda_)
    st.write(f"Mission Success Probability (R): {R:.3f}")
    st.write(f"Failure Avoidance Probability (S): {S:.3f}")
    
    # Objective 2
    st.subheader("Objective 2: Maximize S")
    R_de, S_de, tau_de = objective_2(N, lambda_)
    st.write(f"Mission Success Probability (R): {R_de:.3f}")
    st.write(f"Failure Avoidance Probability (S): {S_de:.3f}")
    st.write(f"Optimal Inspection Time (tau_1): {tau_de[0]:.3f} hr")
    
    # Data for Objective 2 Plot
    lambda_values = [0.15, 0.25, 0.35, 0.45]
    data = {
        'Lambda': lambda_values,
        'R_equals_S': [],
        'S_at_S_max': [],
        'R_at_S_max': [],
        'Tau_at_S_max': []
    }
    for lambda_val in lambda_values:
        R, S = objective_1(lambda_val)
        data['R_equals_S'].append(R)
        R_de, S_de, tau_de = objective_2(N, lambda_val)
        data['S_at_S_max'].append(S_de)
        data['R_at_S_max'].append(R_de)
        data['Tau_at_S_max'].append(tau_de[0])
    
    df = pd.DataFrame(data)
    
    # Plotting Objective 2
    st.subheader("Objective 2 Plot")
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(df['Lambda'], df['R_equals_S'], marker='o', label='Max R when R = S', color='green', linewidth=2)
    ax.plot(df['Lambda'], df['S_at_S_max'], marker='s', label='S at Max S', color='#87CEEB', linewidth=2)
    ax.plot(df['Lambda'], df['R_at_S_max'], marker='^', label='R at Max S', color='orange', linewidth=2)
    ax.set_xlabel('Shock Rate (lambda)')
    ax.set_ylabel('Probability')
    ax.grid(True)
    ax.legend()
    plt.tight_layout()
    st.pyplot(fig)
    
    # Objective 3
    st.subheader("Objective 3: Maximize R s.t. S >= 0.90")
    R_de, S_de, tau_de = objective_3(lambda_)
    st.write(f"Mission Success Probability (R): {R_de:.3f}")
    st.write(f"Failure Avoidance Probability (S): {S_de:.3f}")
    st.write(f"Optimal Inspection Time (tau_1): {tau_de[0]:.3f} hr")
    
    # Data for Objective 3 Plot
    max_r_s90_S = []
    max_r_s90_R = []
    tau1_s90 = []
    for lambda_val in lambda_values:
        R_de, S_de, tau_de = objective_3(lambda_val)
        max_r_s90_S.append(S_de)
        max_r_s90_R.append(R_de)
        tau1_s90.append('NI' if S_de < S_star_90 else f'{tau_de[0]:.2f}')
    
    # Plotting Objective 3
    st.subheader("Objective 3 Plot")
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(lambda_values, max_r_s90_S, label='S Max R (S >= 0.90)', marker='o')
    ax.plot(lambda_values, max_r_s90_R, label='R Max R (S >= 0.90)', marker='^')
    for i, txt in enumerate(tau1_s90):
        if txt != 'NI':
            ax.annotate(f'tau(1)={txt}', (lambda_values[i], max_r_s90_R[i]), 
                        textcoords="offset points", xytext=(0, 10), ha='center')
    ax.set_xlabel('Shock Intensity (lambda)')
    ax.set_ylabel('Probability')
    ax.legend()
    ax.grid(True)
    ax.set_ylim(0, 1.0)
    plt.tight_layout()
    st.pyplot(fig)
    
    # Objective 4
    st.subheader("Objective 4: Maximize R s.t. S >= 0.85")
    R_de, S_de, tau_de = objective_4(lambda_)
    st.write(f"Mission Success Probability (R): {R_de:.3f}")
    st.write(f"Failure Avoidance Probability (S): {S_de:.3f}")
    st.write(f"Optimal Inspection Time (tau_1): {tau_de[0]:.3f} hr")
    
    # Data for Objective 4 Plot
    max_r_s85_S = []
    max_r_s85_R = []
    tau1_s85 = []
    for lambda_val in lambda_values:
        R_de, S_de, tau_de = objective_4(lambda_val)
        max_r_s85_S.append(S_de)
        max_r_s85_R.append(R_de)
        tau1_s85.append('NI' if S_de < S_star_85 else f'{tau_de[0]:.2f}')
    
    # Plotting Objective 4
    st.subheader("Objective 4 Plot")
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(lambda_values, max_r_s85_S, label='S Max R (S >= 0.85)', marker='o')
    ax.plot(lambda_values, max_r_s85_R, label='R Max R (S >= 0.85)', marker='^')
    for i, txt in enumerate(tau1_s85):
        if txt != 'NI':
            ax.annotate(f'tau(1)={txt}', (lambda_values[i], max_r_s85_R[i]), 
                        textcoords="offset points", xytext=(0, 10), ha='center')
    ax.set_xlabel('Shock Intensity (lambda)')
    ax.set_ylabel('Probability')
    ax.legend()
    ax.grid(True)
    ax.set_ylim(0, 1.0)
    plt.tight_layout()
    st.pyplot(fig)

# Add date and time and copyright notice at the bottom
st.write("Last updated: Friday, June 06, 2025, 12:17 PM -03")
st.write("© 2025 All rights reserved.")
